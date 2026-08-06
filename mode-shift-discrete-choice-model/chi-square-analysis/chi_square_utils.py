#!/usr/bin/env python3
"""
chi_square_utils.py
===================
Core statistical analysis engine for Chi-Square tests of independence,
assumptions checking, Monte Carlo exact tests, bias-corrected Cramér's V,
and publication table exports (CSV, Word DOCX, LaTeX).

Author : Majharul Islam
Date   : 2026-08-07
Python : 3.11+
"""

import logging
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
import numpy as np
import pandas as pd
from scipy.stats import chi2_contingency

logger = logging.getLogger("chi_square_utils")


def load_dataset(data_path: str) -> pd.DataFrame:
    """Load Excel dataset."""
    if not os.path.exists(data_path):
        raise FileNotFoundError(f"Dataset not found at: {data_path}")
    logger.info("Loading dataset from: %s", data_path)
    df = pd.read_excel(data_path)
    df.columns = df.columns.str.strip()
    logger.info("Dataset loaded: %d rows × %d columns", len(df), len(df.columns))
    return df


def detect_column(df: pd.DataFrame, keywords: List[str]) -> str:
    """Detect column by keywords."""
    for kw in keywords:
        for col in df.columns:
            if kw.lower() in col.lower():
                logger.info("Detected column '%s' via keyword '%s'", col, kw)
                return col
    raise ValueError(f"Could not detect column matching keywords: {keywords}")


def check_missing_values(df: pd.DataFrame, cols: List[str]) -> pd.DataFrame:
    """Remove missing values for specified columns."""
    df_clean = df.dropna(subset=cols)
    diff = len(df) - len(df_clean)
    if diff > 0:
        logger.warning("Removed %d missing rows in columns %s.", diff, cols)
    else:
        logger.info("No missing values found in specified columns.")
    return df_clean


def build_contingency_table(
    df: pd.DataFrame,
    row_col: str,
    col_col: str,
    row_order: Optional[List[str]] = None,
    col_order: Optional[List[str]] = None,
) -> pd.DataFrame:
    """Build contingency table."""
    ct = pd.crosstab(df[row_col], df[col_col])
    if row_order:
        present_rows = [r for r in row_order if r in ct.index]
        if present_rows:
            ct = ct.reindex(index=present_rows)
    if col_order:
        present_cols = [c for c in col_order if c in ct.columns]
        if present_cols:
            ct = ct.reindex(columns=present_cols)
    ct = ct.fillna(0).astype(int)
    logger.info("Contingency table built: %d rows × %d cols | N = %d", ct.shape[0], ct.shape[1], ct.values.sum())
    return ct


def run_chi_square(observed: pd.DataFrame) -> Tuple[float, float, int, pd.DataFrame]:
    """Run Pearson Chi-square test."""
    chi2, p_val, dof, expected = chi2_contingency(observed.values, correction=False)
    expected_df = pd.DataFrame(expected, index=observed.index, columns=observed.columns)
    logger.info("Chi-square test: χ²=%.4f, df=%d, p=%.6f", chi2, dof, p_val)
    return float(chi2), float(p_val), int(dof), expected_df


def check_chi_square_assumptions(expected_df: pd.DataFrame) -> Dict[str, Any]:
    """Check Chi-square assumptions."""
    exp_vals = expected_df.values.flatten()
    total_cells = len(exp_vals)
    below_5 = int(np.sum(exp_vals < 5))
    below_1 = int(np.sum(exp_vals < 1))
    pct_below_5 = round((below_5 / total_cells) * 100, 1)
    met = (pct_below_5 <= 20.0) and (below_1 == 0)

    recommendation = (
        "Pearson Chi-square (assumptions satisfied)"
        if met
        else "Fisher-Freeman-Halton Exact Test or Monte Carlo simulation (assumptions violated)"
    )

    logger.info("Assumption check [%s]: %d/%d cells (%.1f%%) have E < 5", "SATISFIED" if met else "VIOLATED", below_5, total_cells, pct_below_5)

    return {
        "cells_total": total_cells,
        "cells_below_5": below_5,
        "cells_below_1": below_1,
        "pct_below_5": pct_below_5,
        "assumption_met": met,
        "recommendation": recommendation,
    }


def compute_cramers_v(chi2: float, n: int, r: int, c: int) -> float:
    """Compute bias-corrected Cramér's V (Bergsma, 2013)."""
    phi2 = chi2 / n
    r_corr = r - 1
    c_corr = c - 1
    phi2_corr = max(0, phi2 - ((r_corr * c_corr) / (n - 1)))
    r_adj = r_corr - ((r_corr ** 2) / (n - 1))
    c_adj = c_corr - ((c_corr ** 2) / (n - 1))
    denom = min(max(0, r_adj), max(0, c_adj))
    if denom == 0:
        return 0.0
    v = np.sqrt(phi2_corr / denom)
    logger.info("Cramér's V (bias-corrected) = %.4f", v)
    return round(float(v), 4)


def interpret_cramers_v(v: float) -> str:
    """Interpret Cramér's V magnitude."""
    if v < 0.10:
        return "negligible"
    elif v < 0.20:
        return "small"
    elif v < 0.40:
        return "moderate"
    else:
        return "large"


def make_decision(p_val: float, alpha: float = 0.05) -> str:
    """Generate statistical decision."""
    if p_val < alpha:
        return f"Reject H₀ (p = {p_val:.4f} < α = {alpha:.2f}). There is sufficient evidence to conclude a statistically significant association between the two variables."
    else:
        return f"Fail to reject H₀ (p = {p_val:.4f} ≥ α = {alpha:.2f}). There is insufficient evidence of a statistically significant association between the two variables."


def run_monte_carlo_ffh(observed: pd.DataFrame, n_replications: int = 99999, seed: int = 42) -> float:
    """Monte Carlo Fisher-Freeman-Halton exact test."""
    rng = np.random.default_rng(seed)
    obs_chi2, _, _, _ = chi2_contingency(observed.values, correction=False)
    row_sums = observed.sum(axis=1).values
    col_sums = observed.sum(axis=0).values
    total = observed.values.sum()
    col_probs = col_sums / total

    exceed_count = 0
    for _ in range(n_replications):
        sim_table = np.zeros(observed.shape, dtype=int)
        for r_idx, r_sum in enumerate(row_sums):
            sim_table[r_idx] = rng.multinomial(r_sum, col_probs)
        try:
            sim_chi2, _, _, _ = chi2_contingency(sim_table, correction=False)
            if sim_chi2 >= obs_chi2:
                exceed_count += 1
        except Exception:
            pass

    p_mc = (exceed_count + 1) / (n_replications + 1)
    logger.info("Monte Carlo FFH test: p = %.6f (%d/%d replications)", p_mc, exceed_count, n_replications)
    return round(float(p_mc), 6)


def run_monte_carlo_permutation(observed: pd.DataFrame, n_replications: int = 99999, seed: int = 42) -> float:
    """Monte Carlo permutation test."""
    rng = np.random.default_rng(seed)
    obs_chi2, _, _, _ = chi2_contingency(observed.values, correction=False)
    rows_flat, cols_flat = [], []
    for r_idx in range(observed.shape[0]):
        for c_idx in range(observed.shape[1]):
            cnt = observed.values[r_idx, c_idx]
            rows_flat.extend([r_idx] * cnt)
            cols_flat.extend([c_idx] * cnt)

    rows_arr = np.array(rows_flat)
    cols_arr = np.array(cols_flat)
    exceed_count = 0

    for _ in range(n_replications):
        shuffled_cols = rng.permutation(cols_arr)
        perm_table, _, _ = np.histogram2d(rows_arr, shuffled_cols, bins=(observed.shape[0], observed.shape[1]))
        try:
            perm_chi2, _, _, _ = chi2_contingency(perm_table, correction=False)
            if perm_chi2 >= obs_chi2:
                exceed_count += 1
        except Exception:
            pass

    p_mc = (exceed_count + 1) / (n_replications + 1)
    logger.info("Monte Carlo permutation test: p = %.6f (%d/%d replications)", p_mc, exceed_count, n_replications)
    return round(float(p_mc), 6)


def export_to_csv(df: pd.DataFrame, filepath: str) -> None:
    """Export DataFrame to CSV."""
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    df.to_csv(filepath, encoding="utf-8")
    logger.info("Exported CSV: %s", filepath)


def export_results_csv(results: Dict[str, Any], filepath: str) -> None:
    """Export results dictionary to CSV."""
    df = pd.DataFrame([results])
    export_to_csv(df, filepath)


def build_publication_table(
    observed: pd.DataFrame,
    expected_df: pd.DataFrame,
    chi2: float,
    p_val: float,
    dof: int,
    cramers_v: float,
    v_label: str,
    decision: str,
    row_var_label: str,
    col_var_label: str,
    n_total: int,
    mc_p: Optional[float] = None,
) -> pd.DataFrame:
    """Build publication table."""
    row_totals = observed.sum(axis=1)
    col_totals = observed.sum(axis=0)

    rows = []
    for r_label in observed.index:
        row_data = {row_var_label: str(r_label)}
        r_tot = int(row_totals[r_label])
        for c_label in observed.columns:
            cnt = int(observed.loc[r_label, c_label])
            pct = (cnt / r_tot * 100) if r_tot > 0 else 0
            row_data[str(c_label)] = f"{cnt} ({pct:.1f}%)"
        row_data["Row Total"] = r_tot
        rows.append(row_data)

    tot_data = {row_var_label: "Column Total"}
    for c_label in observed.columns:
        tot_data[str(c_label)] = int(col_totals[c_label])
    tot_data["Row Total"] = n_total
    rows.append(tot_data)

    pub_df = pd.DataFrame(rows)
    return pub_df


def export_publication_docx(
    observed: pd.DataFrame,
    expected_df: pd.DataFrame,
    chi2: float,
    p_val: float,
    dof: int,
    cramers_v: float,
    v_label: str,
    decision: str,
    row_var_label: str,
    col_var_label: str,
    table_title: str,
    filepath: str,
    mc_p: Optional[float] = None,
) -> None:
    """Export APA-7 Word table."""
    doc = docx.Document()
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(table_title)
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(11)
    r_title.font.bold = True

    pub_df = build_publication_table(observed, expected_df, chi2, p_val, dof, cramers_v, v_label, decision, row_var_label, col_var_label, sum(observed.values.flatten()), mc_p)

    table = doc.add_table(rows=len(pub_df) + 1, cols=len(pub_df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col_name in enumerate(pub_df.columns):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        p = cell.paragraphs[0]
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.bold = True

    for r_idx, row in pub_df.iterrows():
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(9.5)

    def set_cell_border(cell, **kwargs):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/><w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/></w:tcBorders>' % (
            nsdecls('w'),
            kwargs.get('top', 'none'), kwargs.get('top_sz', '4'), kwargs.get('top_color', 'auto'),
            kwargs.get('bottom', 'none'), kwargs.get('bottom_sz', '4'), kwargs.get('bottom_color', 'auto')
        ))
        tcPr.append(tcBorders)

    for cell in table.rows[0].cells:
        set_cell_border(cell, top='single', top_sz='12', bottom='single', bottom_sz='8')
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom='single', bottom_sz='12')

    p_fmt = "< 0.001" if p_val < 0.001 else f"{p_val:.4f}"
    mc_note = f" Monte Carlo p = {mc_p:.6f} (99,999 replications, seed = 42)." if mc_p is not None else ""

    note_p = doc.add_paragraph()
    n_tot = sum(observed.values.flatten())
    r_note = note_p.add_run(
        f"Note. N = {n_tot}. Cells show count (row %). χ²({dof}, N={n_tot}) = {chi2:.4f}, p = {p_fmt}.{mc_note} Cramér's V = {cramers_v:.4f} ({v_label} association). Decision: {decision}"
    )
    r_note.font.name = "Times New Roman"
    r_note.font.size = Pt(9)
    r_note.font.italic = True

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    logger.info("Exported DOCX: %s", filepath)


def export_publication_latex(
    observed: pd.DataFrame,
    expected_df: pd.DataFrame,
    chi2: float,
    p_val: float,
    dof: int,
    cramers_v: float,
    v_label: str,
    decision: str,
    row_var_label: str,
    col_var_label: str,
    table_title: str,
    filepath: str,
    mc_p: Optional[float] = None,
) -> None:
    """Export publication LaTeX booktabs table."""
    n_tot = sum(observed.values.flatten())
    row_totals = observed.sum(axis=1)
    col_totals = observed.sum(axis=0)

    cols = list(observed.columns)
    col_spec = "l" + "r" * (len(cols) + 1)
    p_fmt = "$<0.001$" if p_val < 0.001 else f"{p_val:.4f}"

    mc_note = ""
    if mc_p is not None:
        mc_p_fmt = f"{mc_p:.6f}" if mc_p >= 0.001 else "$<0.001$"
        mc_note = f" Monte Carlo $p = {mc_p_fmt}$ (99{{,}}999 replications, seed = 42)."

    lines = [
        "\\begin{table}[htbp]",
        "  \\centering",
        f"  \\caption{{{table_title}}}",
        f"  \\label{{tab:{re.sub(r'[^a-zA-Z0-9]', '_', table_title.lower())[:40]}}}",
        "  \\small",
        f"  \\begin{{tabular}}{{{col_spec}}}",
        "    \\toprule",
    ]

    hdr_parts = [f"\\textbf{{{row_var_label}}}"] + [f"\\textbf{{{c}}}" for c in cols] + ["\\textbf{Total}"]
    lines.append("    " + " & ".join(hdr_parts) + " \\\\")
    lines.append("    \\midrule")

    for r_label in observed.index:
        r_tot = int(row_totals[r_label])
        r_parts = [str(r_label)]
        for c_label in cols:
            cnt = int(observed.loc[r_label, c_label])
            pct = (cnt / r_tot * 100) if r_tot > 0 else 0
            r_parts.append(f"{cnt} ({pct:.1f}\\%)")
        r_parts.append(str(r_tot))
        lines.append("    " + " & ".join(r_parts) + " \\\\")

    lines.append("    \\midrule")
    tot_parts = ["\\textbf{Column Total}"] + [str(int(col_totals[c])) for c in cols] + [f"\\textbf{{{n_tot}}}"]
    lines.append("    " + " & ".join(tot_parts) + " \\\\")
    lines.append("    \\bottomrule")
    lines.append("  \\end{tabular}")
    lines.append("  \\vspace{2pt}")
    lines.append("  \\begin{flushleft}")
    lines.append("  \\footnotesize")
    lines.append(
        f"  \\textit{{Note.}} $N = {n_tot}$. Cells show count (row\\%). $\\chi^2({dof}, N={n_tot}) = {chi2:.4f}$, $p = {p_fmt}$.{mc_note} Cram\\'er's $V = {cramers_v:.4f}$ ({v_label} association)."
    )
    lines.append("  \\end{flushleft}")
    lines.append("\\end{table}")

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    logger.info("Exported LaTeX: %s", filepath)


def run_full_analysis(
    data_path: str,
    row_keywords: List[str],
    col_keywords: List[str],
    row_var_label: str,
    col_var_label: str,
    table_title: str,
    output_dir: str,
    row_order: Optional[List[str]] = None,
    col_order: Optional[List[str]] = None,
    use_monte_carlo_ffh: bool = False,
    use_monte_carlo_permutation: bool = False,
    n_replications: int = 99999,
    seed: int = 42,
    row_recode: Optional[Dict[Any, str]] = None,
    col_recode: Optional[Dict[Any, str]] = None,
) -> Dict[str, Any]:
    """Run full chi-square analysis wrapper."""
    os.makedirs(output_dir, exist_ok=True)
    df = load_dataset(data_path)
    row_col = detect_column(df, row_keywords)
    col_col = detect_column(df, col_keywords)

    df = check_missing_values(df, [row_col, col_col])

    if row_recode:
        df["_row_recoded"] = df[row_col].map(row_recode)
        df = df.dropna(subset=["_row_recoded"])
        row_col = "_row_recoded"

    if col_recode:
        df["_col_recoded"] = df[col_col].map(col_recode)
        df = df.dropna(subset=["_col_recoded"])
        col_col = "_col_recoded"

    n_total = len(df)
    observed = build_contingency_table(df, row_col, col_col, row_order, col_order)
    chi2, p_val, dof, expected_df = run_chi_square(observed)
    assumptions = check_chi_square_assumptions(expected_df)
    cramers_v = compute_cramers_v(chi2, n_total, observed.shape[0], observed.shape[1])
    v_interp = interpret_cramers_v(cramers_v)

    mc_ffh_p = None
    mc_perm_p = None
    if use_monte_carlo_ffh:
        mc_ffh_p = run_monte_carlo_ffh(observed, n_replications, seed)
    if use_monte_carlo_permutation:
        mc_perm_p = run_monte_carlo_permutation(observed, n_replications, seed)

    primary_p = mc_ffh_p if mc_ffh_p is not None else p_val
    decision = make_decision(primary_p)
    mc_p_for_export = mc_ffh_p if mc_ffh_p is not None else mc_perm_p

    results = {
        "analysis": table_title,
        "row_variable": row_var_label,
        "col_variable": col_var_label,
        "n_total": n_total,
        "chi2_statistic": round(chi2, 4),
        "degrees_of_freedom": dof,
        "p_value": round(p_val, 6),
        "cramers_v": cramers_v,
        "effect_size_label": v_interp,
        "cells_total": assumptions["cells_total"],
        "cells_below_5": assumptions["cells_below_5"],
        "pct_cells_below_5": assumptions["pct_below_5"],
        "assumption_met": assumptions["assumption_met"],
        "test_used": assumptions["recommendation"],
        "decision": decision,
    }
    if mc_ffh_p is not None:
        results["monte_carlo_ffh_p"] = mc_ffh_p
    if mc_perm_p is not None:
        results["monte_carlo_permutation_p"] = mc_perm_p

    export_to_csv(observed, os.path.join(output_dir, "observed_frequency.csv"))
    export_to_csv(expected_df.round(4), os.path.join(output_dir, "expected_frequency.csv"))
    export_results_csv(results, os.path.join(output_dir, "results.csv"))

    pub_table = build_publication_table(
        observed, expected_df, chi2, p_val, dof, cramers_v, v_interp,
        decision, row_var_label, col_var_label, n_total, mc_p_for_export
    )
    export_to_csv(pub_table, os.path.join(output_dir, "publication_table.csv"))

    export_publication_docx(
        observed, expected_df, chi2, p_val, dof, cramers_v, v_interp,
        decision, row_var_label, col_var_label, table_title,
        os.path.join(output_dir, "publication_table.docx"), mc_p_for_export
    )

    export_publication_latex(
        observed, expected_df, chi2, p_val, dof, cramers_v, v_interp,
        decision, row_var_label, col_var_label, table_title,
        os.path.join(output_dir, "publication_table.tex"), mc_p_for_export
    )

    logger.info("All outputs saved to: %s", output_dir)
    return results

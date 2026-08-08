#!/usr/bin/env python3
"""
mann_whitney_utils.py
=====================
Core statistical engine for the Mann-Whitney U reproducibility pipeline.
All calculations are performed from scratch to ensure full transparency
and auditability. Library results (scipy) are used only for cross-validation.

Author : Majharul Islam
Date   : 2026-08-09
Python : 3.10+
"""

import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
from scipy.stats import mannwhitneyu, norm, rankdata

logger = logging.getLogger("mann_whitney_utils")

# ============================================================
# VARIABLE MAPPING: Raw Excel column → short analysis name
# ============================================================
RAW_COLUMN_MAP = {
    "security_harassment": (
        "How safe do you feel regarding harassment, pickpocketing, "
        "or personal security?"
    ),
    "reliability": "How reliable is your current mode? (Does it arrive on time?)",
    "road_accidents": (
        "How safe do you feel regarding road accidents on this mode?"
    ),
    "comfort": (
        "How would you rate the physical comfort (Seating space, AC, Noise) "
        "of your current mode?"
    ),
    "crowding": "How crowded is the vehicle usually?",
}

GENDER_RAW_COLUMN = "Gender"  # after stripping whitespace

ANALYSIS_LABELS = {
    "security_harassment": "Safety: Harassment/Security",
    "reliability": "Reliability",
    "road_accidents": "Safety: Road Accidents",
    "comfort": "Comfort",
    "crowding": "Crowding",
}

# Reported thesis values (Table 4.4) — NEVER used in calculations
REPORTED_THESIS = {
    "security_harassment": {"U": 9399.50, "p": 0.015},
    "reliability": {"U": 12340.50, "p": 0.142},
    "road_accidents": {"U": 11610.00, "p": 0.619},
    "comfort": {"U": 10999.00, "p": 0.753},
    "crowding": {"U": 10837.00, "p": 0.598},
}


# ============================================================
# DATA LOADING
# ============================================================
def load_raw_data(data_path: str) -> pd.DataFrame:
    """Load raw Excel dataset and normalise column names.

    Column names are stripped of leading/trailing whitespace.
    No rows are dropped at this stage.
    """
    if not os.path.exists(data_path):
        raise FileNotFoundError(f"Dataset not found: {data_path}")

    df = pd.read_excel(data_path)
    df.columns = df.columns.str.strip()
    logger.info("Loaded %d rows × %d columns from %s", len(df), len(df.columns), data_path)
    return df


def detect_column(df: pd.DataFrame, short_name: str) -> str:
    """Find the raw column matching a short analysis name.

    Uses substring matching against RAW_COLUMN_MAP targets.
    """
    target_fragment = RAW_COLUMN_MAP[short_name]
    for col in df.columns:
        col_clean = col.strip()
        if target_fragment.lower() in col_clean.lower():
            return col
    raise ValueError(
        f"Cannot find column for '{short_name}'. "
        f"Searched for: '{target_fragment}'"
    )


def extract_gender(df: pd.DataFrame) -> pd.Series:
    """Extract and validate the Gender column.

    Returns the raw series without modification.
    Gender values are expected to be text: 'Male' / 'Female'.
    """
    gender_col = None
    for col in df.columns:
        if "gender" in col.strip().lower():
            gender_col = col
            break
    if gender_col is None:
        raise ValueError("Gender column not found in dataset")

    gender = df[gender_col].copy()
    return gender


# ============================================================
# DATA VALIDATION (STEP 2)
# ============================================================
def validate_data(
    df: pd.DataFrame,
    outcome_cols: Dict[str, str],
    gender_series: pd.Series,
) -> Dict[str, Any]:
    """Comprehensive data validation report.

    Parameters
    ----------
    df : raw dataframe
    outcome_cols : {short_name: actual_column_name}
    gender_series : extracted gender series

    Returns
    -------
    dict with validation results
    """
    report: Dict[str, Any] = {}

    # Global checks
    report["total_observations"] = len(df)
    report["duplicate_rows"] = int(df.duplicated().sum())

    # Gender checks
    gender_clean = gender_series.astype(str).str.strip()
    report["missing_gender"] = int(gender_series.isna().sum())
    report["gender_value_counts"] = gender_clean.value_counts().to_dict()
    valid_genders = {"Male", "Female"}
    invalid_mask = ~gender_clean.isin(valid_genders) & gender_series.notna()
    report["invalid_gender_codes"] = int(invalid_mask.sum())
    if report["invalid_gender_codes"] > 0:
        report["invalid_gender_values"] = (
            gender_clean[invalid_mask].unique().tolist()
        )

    n_male = int((gender_clean == "Male").sum())
    n_female = int((gender_clean == "Female").sum())
    report["n_male"] = n_male
    report["n_female"] = n_female
    report["total_valid_gender"] = n_male + n_female

    # Per-outcome checks
    outcome_reports = {}
    for short_name, col_name in outcome_cols.items():
        vals = df[col_name]
        orep: Dict[str, Any] = {}
        orep["column_name"] = col_name
        orep["N"] = len(vals)
        orep["missing"] = int(vals.isna().sum())
        orep["valid_N"] = int(vals.notna().sum())

        valid_vals = vals.dropna()
        if len(valid_vals) > 0:
            numeric_vals = pd.to_numeric(valid_vals, errors="coerce")
            orep["non_numeric_count"] = int(numeric_vals.isna().sum())
            numeric_clean = numeric_vals.dropna()
            if len(numeric_clean) > 0:
                orep["minimum"] = float(numeric_clean.min())
                orep["maximum"] = float(numeric_clean.max())
                orep["mean"] = round(float(numeric_clean.mean()), 4)
                orep["median"] = float(numeric_clean.median())
                orep["sd"] = round(float(numeric_clean.std(ddof=1)), 4)

                # Check for out-of-range Likert values
                outside_range = numeric_clean[
                    (numeric_clean < 1) | (numeric_clean > 5)
                ]
                orep["values_outside_1_5"] = int(len(outside_range))
                if orep["values_outside_1_5"] > 0:
                    orep["out_of_range_values"] = outside_range.tolist()

                # Frequency distribution
                freq = {}
                for v in range(1, 6):
                    freq[v] = int((numeric_clean == v).sum())
                orep["frequency_1_to_5"] = freq
                orep["frequency_sum"] = sum(freq.values())

        outcome_reports[short_name] = orep

    report["outcome_variables"] = outcome_reports
    return report


def format_validation_report(report: Dict[str, Any]) -> str:
    """Format validation report as readable text."""
    lines = []
    lines.append("=" * 70)
    lines.append("DATA VALIDATION REPORT")
    lines.append("=" * 70)
    lines.append(f"Total observations: {report['total_observations']}")
    lines.append(f"Duplicate rows: {report['duplicate_rows']}")
    lines.append(f"Missing Gender: {report['missing_gender']}")
    lines.append(f"Invalid Gender codes: {report['invalid_gender_codes']}")
    if report.get("invalid_gender_values"):
        lines.append(f"  Invalid values: {report['invalid_gender_values']}")
    lines.append(f"N Male: {report['n_male']}")
    lines.append(f"N Female: {report['n_female']}")
    lines.append(f"Total valid Gender: {report['total_valid_gender']}")
    lines.append(f"Gender distribution: {report['gender_value_counts']}")

    for short_name, orep in report["outcome_variables"].items():
        lines.append("")
        lines.append(f"--- {ANALYSIS_LABELS.get(short_name, short_name)} ---")
        lines.append(f"  Column: {orep.get('column_name', 'N/A')}")
        lines.append(f"  N: {orep.get('N', 'N/A')}")
        lines.append(f"  Missing: {orep.get('missing', 'N/A')}")
        lines.append(f"  Valid N: {orep.get('valid_N', 'N/A')}")
        if "minimum" in orep:
            lines.append(f"  Min: {orep['minimum']}")
            lines.append(f"  Max: {orep['maximum']}")
            lines.append(f"  Mean: {orep['mean']}")
            lines.append(f"  Median: {orep['median']}")
            lines.append(f"  SD: {orep['sd']}")
        if "values_outside_1_5" in orep:
            lines.append(f"  Values outside 1–5: {orep['values_outside_1_5']}")
        if "frequency_1_to_5" in orep:
            freq = orep["frequency_1_to_5"]
            lines.append(f"  Freq: 1={freq[1]}, 2={freq[2]}, 3={freq[3]}, "
                         f"4={freq[4]}, 5={freq[5]}  (sum={orep['frequency_sum']})")

    lines.append("")
    return "\n".join(lines)


# ============================================================
# MANN-WHITNEY U CALCULATION (STEPS 3-8)
# ============================================================
def compute_mann_whitney_full(
    male_values: np.ndarray,
    female_values: np.ndarray,
    variable_name: str = "",
) -> Dict[str, Any]:
    """Complete Mann-Whitney U calculation from scratch.

    Parameters
    ----------
    male_values : array of ordinal ratings for males
    female_values : array of ordinal ratings for females
    variable_name : label for reporting

    Returns
    -------
    dict with ALL intermediate and final statistics
    """
    result: Dict[str, Any] = {"variable": variable_name}

    n1 = len(male_values)  # Male
    n2 = len(female_values)  # Female
    N = n1 + n2
    result["n1_male"] = n1
    result["n2_female"] = n2
    result["N"] = N

    # Descriptive statistics
    result["male_mean"] = round(float(np.mean(male_values)), 4)
    result["male_median"] = float(np.median(male_values))
    result["male_sd"] = round(float(np.std(male_values, ddof=1)), 4)
    result["female_mean"] = round(float(np.mean(female_values)), 4)
    result["female_median"] = float(np.median(female_values))
    result["female_sd"] = round(float(np.std(female_values, ddof=1)), 4)

    # Male mean rank and female mean rank
    result["male_mean_rank"] = None  # set below after ranking
    result["female_mean_rank"] = None

    # ---- A. Combined ranks ----
    # Pool all observations, rank with average ties
    combined = np.concatenate([male_values, female_values])
    ranks = rankdata(combined, method="average")

    R1 = float(np.sum(ranks[:n1]))  # Male rank sum
    R2 = float(np.sum(ranks[n1:]))  # Female rank sum
    result["R1_male"] = R1
    result["R2_female"] = R2

    result["male_mean_rank"] = round(R1 / n1, 4)
    result["female_mean_rank"] = round(R2 / n2, 4)

    # Identity check: R1 + R2 = N(N+1)/2
    expected_rank_sum = N * (N + 1) / 2
    result["R1_plus_R2"] = R1 + R2
    result["expected_rank_sum"] = expected_rank_sum
    result["rank_sum_identity_ok"] = abs(R1 + R2 - expected_rank_sum) < 1e-6

    if not result["rank_sum_identity_ok"]:
        logger.error(
            "RANK SUM IDENTITY FAILED for %s: R1+R2=%.1f != N(N+1)/2=%.1f",
            variable_name, R1 + R2, expected_rank_sum,
        )

    # ---- B. U1 (Male) ----
    U1 = n1 * n2 + n1 * (n1 + 1) / 2 - R1
    result["U1_male"] = U1

    # ---- C. U2 (Female) ----
    U2 = n1 * n2 + n2 * (n2 + 1) / 2 - R2
    result["U2_female"] = U2

    # ---- D. U identity check ----
    result["U1_plus_U2"] = U1 + U2
    result["n1_times_n2"] = n1 * n2
    result["U_identity_ok"] = abs(U1 + U2 - n1 * n2) < 1e-6

    if not result["U_identity_ok"]:
        logger.error(
            "U IDENTITY FAILED for %s: U1+U2=%.1f != n1*n2=%d",
            variable_name, U1 + U2, n1 * n2,
        )

    # ---- E. Reported U = min(U1, U2) ----
    U = min(U1, U2)
    result["U_min"] = U
    result["U_direction"] = "U1 (Male)" if U == U1 else "U2 (Female)"

    # ---- STEP 4: Expected U ----
    mu_U = n1 * n2 / 2
    result["mu_U"] = mu_U

    # ---- STEP 5: Tie correction ----
    unique_vals, counts = np.unique(combined, return_counts=True)
    tie_groups = []
    for val, cnt in zip(unique_vals, counts):
        tie_groups.append({
            "value": float(val),
            "count": int(cnt),
            "t_cubed_minus_t": int(cnt ** 3 - cnt),
        })
    result["tie_groups"] = tie_groups

    tie_sum = float(np.sum(counts ** 3 - counts))
    result["tie_sum_t3_minus_t"] = tie_sum

    # Verify tie counts sum to N
    total_tie_counts = int(np.sum(counts))
    result["tie_counts_sum"] = total_tie_counts
    result["tie_counts_sum_ok"] = total_tie_counts == N

    # Standard tie-corrected variance:
    # Var(U) = (n1*n2/12) * [(N+1) - Σ(ti³-ti)/(N*(N-1))]
    var_U = (n1 * n2 / 12) * ((N + 1) - tie_sum / (N * (N - 1)))
    sigma_U = np.sqrt(var_U)
    result["var_U"] = var_U
    result["sigma_U"] = sigma_U

    # Also compute variance without tie correction for comparison
    var_U_no_ties = (n1 * n2 / 12) * (N + 1)
    sigma_U_no_ties = np.sqrt(var_U_no_ties)
    result["var_U_no_ties"] = var_U_no_ties
    result["sigma_U_no_ties"] = sigma_U_no_ties

    # ---- STEP 6: Z calculation ----
    # Without continuity correction
    Z_no_CC = (U - mu_U) / sigma_U
    result["Z_no_CC"] = Z_no_CC

    # With continuity correction (shift U 0.5 toward the mean)
    if U < mu_U:
        cc = 0.5
    elif U > mu_U:
        cc = -0.5
    else:
        cc = 0.0
    Z_CC = (U - mu_U + cc) / sigma_U
    result["Z_CC"] = Z_CC
    result["continuity_correction"] = cc

    # ---- STEP 7: P-values ----
    # Manual two-tailed p from normal distribution
    p_no_CC = float(2 * norm.sf(abs(Z_no_CC)))
    p_CC = float(2 * norm.sf(abs(Z_CC)))
    result["p_no_CC_manual"] = p_no_CC
    result["p_CC_manual"] = p_CC

    try:
        res_scipy_cc = mannwhitneyu(
            male_values, female_values,
            alternative="two-sided",
            method="asymptotic",
            use_continuity=True,
        )
        # SciPy returns R1 - n1(n1+1)/2. Converted to standard U1 (n1*n2 + n1(n1+1)/2 - R1):
        result["scipy_U1"] = float(n1 * n2 - res_scipy_cc.statistic)
        result["scipy_p_CC"] = float(res_scipy_cc.pvalue)
    except Exception as e:
        logger.warning("scipy mannwhitneyu (CC) failed: %s", e)
        result["scipy_U1"] = None
        result["scipy_p_CC"] = None

    try:
        res_scipy_nocc = mannwhitneyu(
            male_values, female_values,
            alternative="two-sided",
            method="asymptotic",
            use_continuity=False,
        )
        result["scipy_p_no_CC"] = float(res_scipy_nocc.pvalue)
    except Exception as e:
        logger.warning("scipy mannwhitneyu (no CC) failed: %s", e)
        result["scipy_p_no_CC"] = None

    # ---- STEP 8: Effect size ----
    # r = |Z| / sqrt(N)  (using Z without CC as is standard)
    r_effect = abs(Z_no_CC) / np.sqrt(N)
    result["effect_size_r"] = round(float(r_effect), 4)
    result["effect_size_r_interpretation"] = interpret_effect_size_r(r_effect)

    # Rank-biserial correlation: rB = 1 - (2*U)/(n1*n2)
    # where U = min(U1, U2) gives rB for the direction of the effect
    r_biserial = 1 - (2 * U) / (n1 * n2)
    result["rank_biserial_r"] = round(float(r_biserial), 4)

    # Also compute using the formula: rB = (R1/n1 - R2/n2) / N
    # This gives a signed value showing direction
    r_biserial_signed = (2 * U1) / (n1 * n2) - 1
    result["rank_biserial_signed"] = round(float(r_biserial_signed), 4)

    # ---- Statistical decision ----
    alpha = 0.05
    # Primary decision uses Z without CC (more common in literature)
    result["alpha"] = alpha
    result["significant_no_CC"] = p_no_CC < alpha
    result["significant_CC"] = p_CC < alpha

    if p_no_CC < alpha:
        result["decision"] = (
            f"Reject H₀ (p = {p_no_CC:.6f} < α = {alpha}). "
            f"Statistically significant difference between Male and Female."
        )
    else:
        result["decision"] = (
            f"Fail to reject H₀ (p = {p_no_CC:.6f} ≥ α = {alpha}). "
            f"No statistically significant difference between Male and Female."
        )

    return result


def interpret_effect_size_r(r: float) -> str:
    """Interpret effect size r using Cohen's benchmarks."""
    r = abs(r)
    if r < 0.10:
        return "negligible"
    elif r < 0.30:
        return "small"
    elif r < 0.50:
        return "medium"
    else:
        return "large"


# ============================================================
# MULTIPLE TESTING CORRECTION (STEP 9)
# ============================================================
def multiple_testing_correction(
    p_values: Dict[str, float],
) -> Dict[str, Dict[str, float]]:
    """Apply Bonferroni and Holm corrections to p-values.

    Parameters
    ----------
    p_values : {variable_name: raw_p_value}

    Returns
    -------
    {variable_name: {raw_p, bonferroni_p, holm_p}}
    """
    k = len(p_values)
    names = list(p_values.keys())
    raw = np.array([p_values[n] for n in names])

    results = {}

    # Bonferroni: p_adj = min(p * k, 1.0)
    bonferroni = np.minimum(raw * k, 1.0)

    # Holm step-down procedure
    sorted_indices = np.argsort(raw)
    holm = np.ones(k)
    for rank_pos, idx in enumerate(sorted_indices):
        holm[idx] = raw[idx] * (k - rank_pos)
    # Enforce monotonicity: holm[i] >= holm[j] if p[i] > p[j] in sorted order
    holm_sorted = holm[sorted_indices]
    for i in range(1, k):
        if holm_sorted[i] < holm_sorted[i - 1]:
            holm_sorted[i] = holm_sorted[i - 1]
    holm[sorted_indices] = holm_sorted
    holm = np.minimum(holm, 1.0)

    for i, name in enumerate(names):
        results[name] = {
            "raw_p": float(raw[i]),
            "bonferroni_p": float(bonferroni[i]),
            "holm_p": float(holm[i]),
            "significant_raw": bool(raw[i] < 0.05),
            "significant_bonferroni": bool(bonferroni[i] < 0.05),
            "significant_holm": bool(holm[i] < 0.05),
        }

    return results


# ============================================================
# CROSS-SOFTWARE COMPARISON TABLE (STEP 11)
# ============================================================
def build_comparison_row(
    variable: str,
    python_result: Dict[str, Any],
    r_result: Optional[Dict[str, float]] = None,
) -> Dict[str, Any]:
    """Build a cross-software comparison row for one variable."""
    row = {
        "Variable": ANALYSIS_LABELS.get(variable, variable),
        "Python_U": python_result["U_min"],
        "Python_U1": python_result["U1_male"],
        "Python_U2": python_result["U2_female"],
        "Python_p_no_CC": python_result["p_no_CC_manual"],
        "Python_p_CC": python_result["p_CC_manual"],
        "SciPy_p_CC": python_result.get("scipy_p_CC"),
    }

    if r_result:
        row["R_U"] = r_result.get("U")
        row["R_p"] = r_result.get("p")
    else:
        row["R_U"] = "—"
        row["R_p"] = "—"

    # SPSS is manual-execution, placeholder
    row["SPSS_U"] = "See SPSS output"
    row["SPSS_p"] = "See SPSS output"

    return row


# ============================================================
# REPORTED vs REPRODUCED (STEP 12)
# ============================================================
def build_reported_vs_reproduced(
    results: Dict[str, Dict[str, Any]],
) -> pd.DataFrame:
    """Compare thesis-reported values against computed values.

    Uses REPORTED_THESIS (never modifies calculations to match).
    """
    rows = []
    for var_name, computed in results.items():
        reported = REPORTED_THESIS.get(var_name, {})
        rep_U = reported.get("U")
        rep_p = reported.get("p")
        calc_U = computed["U_min"]
        calc_U1 = computed["U1_male"]
        calc_U2 = computed["U2_female"]
        calc_p = computed["p_no_CC_manual"]

        U_diff = abs(calc_U - rep_U) if rep_U is not None else None
        p_diff = abs(calc_p - rep_p) if rep_p is not None else None

        # Determine status
        if rep_U is not None and rep_p is not None:
            if abs(calc_U - rep_U) < 0.01 and p_diff < 0.001:
                status = "EXACT MATCH"
            elif (abs(calc_U1 - rep_U) < 0.01 or abs(calc_U2 - rep_U) < 0.01) and p_diff < 0.001:
                status = "METHODOLOGICAL DISCREPANCY (Thesis reported U1 instead of U_min)"
            elif p_diff < 0.005:
                status = "MINOR SOFTWARE DIFFERENCE"
            else:
                status = "DATA DISCREPANCY"
        else:
            status = "NOT REPRODUCIBLE"

        rows.append({
            "Variable": ANALYSIS_LABELS.get(var_name, var_name),
            "Reported_U": rep_U,
            "Calculated_U": calc_U,
            "U_difference": U_diff,
            "Reported_p": rep_p,
            "Calculated_p_no_CC": calc_p,
            "Calculated_p_CC": computed["p_CC_manual"],
            "p_difference_no_CC": p_diff,
            "Status": status,
        })

    return pd.DataFrame(rows)


# ============================================================
# DESCRIPTIVE STATISTICS TABLE
# ============================================================
def build_descriptive_table(
    df: pd.DataFrame,
    outcome_cols: Dict[str, str],
    gender_series: pd.Series,
) -> pd.DataFrame:
    """Build descriptive statistics table per variable per gender."""
    rows = []
    gender_clean = gender_series.astype(str).str.strip()

    for short_name, col_name in outcome_cols.items():
        for gender_label in ["Male", "Female"]:
            mask = gender_clean == gender_label
            vals = pd.to_numeric(
                df.loc[mask, col_name], errors="coerce"
            ).dropna()

            row = {
                "Variable": ANALYSIS_LABELS.get(short_name, short_name),
                "Gender": gender_label,
                "N": int(len(vals)),
                "Mean": round(float(vals.mean()), 4) if len(vals) > 0 else None,
                "Median": float(vals.median()) if len(vals) > 0 else None,
                "SD": round(float(vals.std(ddof=1)), 4) if len(vals) > 1 else None,
                "Min": float(vals.min()) if len(vals) > 0 else None,
                "Max": float(vals.max()) if len(vals) > 0 else None,
            }

            # Frequency distribution
            if len(vals) > 0:
                for v in range(1, 6):
                    row[f"Freq_{v}"] = int((vals == v).sum())
            else:
                for v in range(1, 6):
                    row[f"Freq_{v}"] = 0

            rows.append(row)

    return pd.DataFrame(rows)


# ============================================================
# TIE CORRECTION DETAILS TABLE
# ============================================================
def build_tie_correction_table(
    results: Dict[str, Dict[str, Any]],
) -> pd.DataFrame:
    """Build detailed tie correction table from results."""
    rows = []
    for var_name, res in results.items():
        for tg in res.get("tie_groups", []):
            rows.append({
                "Variable": ANALYSIS_LABELS.get(var_name, var_name),
                "Likert_Value": int(tg["value"]),
                "Count_t": tg["count"],
                "t_cubed_minus_t": tg["t_cubed_minus_t"],
            })
        # Summary row
        rows.append({
            "Variable": ANALYSIS_LABELS.get(var_name, var_name),
            "Likert_Value": "TOTAL",
            "Count_t": res["tie_counts_sum"],
            "t_cubed_minus_t": int(res["tie_sum_t3_minus_t"]),
        })

    return pd.DataFrame(rows)


# ============================================================
# EXPORT FUNCTIONS
# ============================================================
def export_full_results_csv(
    results: Dict[str, Dict[str, Any]],
    filepath: str,
) -> None:
    """Export full results (all intermediates) to CSV."""
    rows = []
    for var_name, res in results.items():
        row = {
            "Variable": ANALYSIS_LABELS.get(var_name, var_name),
            "n1_male": res["n1_male"],
            "n2_female": res["n2_female"],
            "N": res["N"],
            "male_mean": res["male_mean"],
            "male_median": res["male_median"],
            "male_sd": res["male_sd"],
            "male_mean_rank": res["male_mean_rank"],
            "female_mean": res["female_mean"],
            "female_median": res["female_median"],
            "female_sd": res["female_sd"],
            "female_mean_rank": res["female_mean_rank"],
            "R1_male": res["R1_male"],
            "R2_female": res["R2_female"],
            "R1_plus_R2": res["R1_plus_R2"],
            "expected_rank_sum": res["expected_rank_sum"],
            "rank_sum_identity_ok": res["rank_sum_identity_ok"],
            "U1_male": res["U1_male"],
            "U2_female": res["U2_female"],
            "U_min": res["U_min"],
            "U_direction": res["U_direction"],
            "U1_plus_U2": res["U1_plus_U2"],
            "n1_times_n2": res["n1_times_n2"],
            "U_identity_ok": res["U_identity_ok"],
            "mu_U": res["mu_U"],
            "tie_sum_t3_minus_t": res["tie_sum_t3_minus_t"],
            "var_U": res["var_U"],
            "sigma_U": res["sigma_U"],
            "var_U_no_ties": res["var_U_no_ties"],
            "sigma_U_no_ties": res["sigma_U_no_ties"],
            "Z_no_CC": res["Z_no_CC"],
            "Z_CC": res["Z_CC"],
            "continuity_correction": res["continuity_correction"],
            "p_no_CC_manual": res["p_no_CC_manual"],
            "p_CC_manual": res["p_CC_manual"],
            "scipy_U1": res.get("scipy_U1"),
            "scipy_p_CC": res.get("scipy_p_CC"),
            "scipy_p_no_CC": res.get("scipy_p_no_CC"),
            "effect_size_r": res["effect_size_r"],
            "effect_size_interpretation": res["effect_size_r_interpretation"],
            "rank_biserial_r": res["rank_biserial_r"],
            "rank_biserial_signed": res["rank_biserial_signed"],
            "significant_no_CC": res["significant_no_CC"],
            "significant_CC": res["significant_CC"],
            "decision": res["decision"],
        }
        rows.append(row)

    df = pd.DataFrame(rows)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    df.to_csv(filepath, index=False)
    logger.info("Exported full results: %s", filepath)


def export_table_4_4_csv(
    results: Dict[str, Dict[str, Any]],
    filepath: str,
) -> None:
    """Export Table 4.4 format CSV."""
    rows = []
    for var_name, res in results.items():
        rows.append({
            "Variable": ANALYSIS_LABELS.get(var_name, var_name),
            "Male_N": res["n1_male"],
            "Female_N": res["n2_female"],
            "Male_Mean_Rank": res["male_mean_rank"],
            "Female_Mean_Rank": res["female_mean_rank"],
            "U_statistic": res["U_min"],
            "Z": round(res["Z_no_CC"], 4),
            "p_value": round(res["p_no_CC_manual"], 6),
            "Effect_size_r": res["effect_size_r"],
            "Significant": res["significant_no_CC"],
        })
    df = pd.DataFrame(rows)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    df.to_csv(filepath, index=False)
    logger.info("Exported Table 4.4: %s", filepath)


def export_table_4_4_xlsx(
    results: Dict[str, Dict[str, Any]],
    filepath: str,
) -> None:
    """Export publication-ready Table 4.4 as Excel."""
    rows = []
    for var_name, res in results.items():
        p_val = res["p_no_CC_manual"]
        rows.append({
            "Travel Experience Variable": ANALYSIS_LABELS.get(var_name, var_name),
            "Male (N)": res["n1_male"],
            "Female (N)": res["n2_female"],
            "Male Mean Rank": res["male_mean_rank"],
            "Female Mean Rank": res["female_mean_rank"],
            "Mann-Whitney U": res["U_min"],
            "Z-statistic": round(res["Z_no_CC"], 4),
            "p-value (two-tailed)": f"{p_val:.4f}" if p_val >= 0.001 else "< 0.001",
            "Effect Size (r)": res["effect_size_r"],
            "Interpretation": res["effect_size_r_interpretation"],
            "Decision (α = 0.05)": (
                "Significant*" if res["significant_no_CC"] else "Not significant"
            ),
        })
    df = pd.DataFrame(rows)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    df.to_excel(filepath, index=False, sheet_name="Table 4.4")
    logger.info("Exported Table 4.4 XLSX: %s", filepath)


# ============================================================
# FIGURE GENERATION (STEP 15)
# ============================================================
def generate_all_figures(
    df: pd.DataFrame,
    outcome_cols: Dict[str, str],
    gender_series: pd.Series,
    results: Dict[str, Dict[str, Any]],
    output_dir: str,
) -> List[str]:
    """Generate all figures for the analysis.

    Returns list of generated file paths.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.ticker import MaxNLocator

    os.makedirs(output_dir, exist_ok=True)
    gender_clean = gender_series.astype(str).str.strip()
    generated = []

    # Color palette
    male_color = "#2196F3"
    female_color = "#E91E63"

    # ---- Figure 1: Stacked/grouped bar charts for each variable ----
    for short_name, col_name in outcome_cols.items():
        label = ANALYSIS_LABELS.get(short_name, short_name)
        fig, ax = plt.subplots(figsize=(8, 5))

        male_vals = pd.to_numeric(
            df.loc[gender_clean == "Male", col_name], errors="coerce"
        ).dropna()
        female_vals = pd.to_numeric(
            df.loc[gender_clean == "Female", col_name], errors="coerce"
        ).dropna()

        x = np.arange(1, 6)
        width = 0.35

        male_freq = [(male_vals == v).sum() for v in x]
        female_freq = [(female_vals == v).sum() for v in x]

        bars1 = ax.bar(x - width / 2, male_freq, width, label="Male",
                       color=male_color, alpha=0.85, edgecolor="white")
        bars2 = ax.bar(x + width / 2, female_freq, width, label="Female",
                       color=female_color, alpha=0.85, edgecolor="white")

        # Add count labels
        for bar in bars1:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width() / 2, h + 0.5,
                        str(int(h)), ha="center", va="bottom", fontsize=8)
        for bar in bars2:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width() / 2, h + 0.5,
                        str(int(h)), ha="center", va="bottom", fontsize=8)

        ax.set_xlabel("Likert Scale Rating", fontsize=11)
        ax.set_ylabel("Frequency", fontsize=11)
        ax.set_title(f"{label}\nMale vs Female Distribution", fontsize=13,
                     fontweight="bold")
        ax.set_xticks(x)
        ax.legend()
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        # Add test result annotation
        res = results.get(short_name, {})
        p_val = res.get("p_no_CC_manual", None)
        u_val = res.get("U_min", None)
        if p_val is not None and u_val is not None:
            sig = "p < 0.05 *" if p_val < 0.05 else f"p = {p_val:.3f}"
            ax.text(0.98, 0.95,
                    f"U = {u_val:.1f}\n{sig}",
                    transform=ax.transAxes, ha="right", va="top",
                    fontsize=9, bbox=dict(boxstyle="round,pad=0.3",
                                          facecolor="lightyellow",
                                          edgecolor="gray", alpha=0.8))

        plt.tight_layout()
        fpath = os.path.join(output_dir, f"fig_{short_name}_distribution.png")
        fig.savefig(fpath, dpi=150, bbox_inches="tight")
        plt.close(fig)
        generated.append(fpath)

    # ---- Figure 2: Box plots ----
    fig, axes = plt.subplots(1, 5, figsize=(20, 5), sharey=True)
    for idx, (short_name, col_name) in enumerate(outcome_cols.items()):
        ax = axes[idx]
        label = ANALYSIS_LABELS.get(short_name, short_name)

        male_vals = pd.to_numeric(
            df.loc[gender_clean == "Male", col_name], errors="coerce"
        ).dropna().values
        female_vals = pd.to_numeric(
            df.loc[gender_clean == "Female", col_name], errors="coerce"
        ).dropna().values

        bp = ax.boxplot(
            [male_vals, female_vals],
            tick_labels=["Male", "Female"],
            patch_artist=True,
            widths=0.6,
            medianprops=dict(color="black", linewidth=2),
        )
        bp["boxes"][0].set_facecolor(male_color)
        bp["boxes"][0].set_alpha(0.6)
        bp["boxes"][1].set_facecolor(female_color)
        bp["boxes"][1].set_alpha(0.6)

        # Jitter overlay
        np.random.seed(42)
        for i, (vals, color) in enumerate(
            [(male_vals, male_color), (female_vals, female_color)]
        ):
            jitter = np.random.normal(0, 0.04, size=len(vals))
            ax.scatter(
                np.full_like(vals, i + 1, dtype=float) + jitter,
                vals, alpha=0.3, s=8, color=color, zorder=3,
            )

        ax.set_title(label, fontsize=10, fontweight="bold")
        ax.set_yticks([1, 2, 3, 4, 5])
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        # Significance marker
        res = results.get(short_name, {})
        if res.get("significant_no_CC"):
            ax.text(0.5, 1.02, "★", transform=ax.transAxes,
                    ha="center", fontsize=14, color="red")

    fig.suptitle("Mann-Whitney U: Male vs Female Comparison (Box Plots)",
                 fontsize=14, fontweight="bold", y=1.02)
    plt.tight_layout()
    fpath = os.path.join(output_dir, "fig_boxplots_all_variables.png")
    fig.savefig(fpath, dpi=150, bbox_inches="tight")
    plt.close(fig)
    generated.append(fpath)

    # ---- Figure 3: Effect size forest plot ----
    fig, ax = plt.subplots(figsize=(8, 5))
    var_labels = []
    effect_sizes = []
    for short_name in outcome_cols:
        res = results.get(short_name, {})
        var_labels.append(ANALYSIS_LABELS.get(short_name, short_name))
        effect_sizes.append(res.get("effect_size_r", 0))

    y_pos = np.arange(len(var_labels))
    colors = ["#E91E63" if e >= 0.3 else "#FF9800" if e >= 0.1 else "#4CAF50"
              for e in effect_sizes]

    ax.barh(y_pos, effect_sizes, color=colors, alpha=0.8, edgecolor="white",
            height=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(var_labels)
    ax.set_xlabel("Effect Size (r = |Z| / √N)", fontsize=11)
    ax.set_title("Effect Size Summary", fontsize=13, fontweight="bold")

    # Cohen's benchmarks
    ax.axvline(x=0.1, color="gray", linestyle="--", alpha=0.5, linewidth=0.8)
    ax.axvline(x=0.3, color="gray", linestyle="--", alpha=0.5, linewidth=0.8)
    ax.axvline(x=0.5, color="gray", linestyle="--", alpha=0.5, linewidth=0.8)
    ax.text(0.1, len(var_labels) - 0.3, "small", fontsize=8, color="gray")
    ax.text(0.3, len(var_labels) - 0.3, "medium", fontsize=8, color="gray")
    ax.text(0.5, len(var_labels) - 0.3, "large", fontsize=8, color="gray")

    for i, v in enumerate(effect_sizes):
        ax.text(v + 0.005, i, f"{v:.3f}", va="center", fontsize=9)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.invert_yaxis()
    plt.tight_layout()
    fpath = os.path.join(output_dir, "fig_effect_size_forest.png")
    fig.savefig(fpath, dpi=150, bbox_inches="tight")
    plt.close(fig)
    generated.append(fpath)

    # ---- Figure 4: Mean rank comparison ----
    fig, ax = plt.subplots(figsize=(10, 5))
    var_labels = []
    male_ranks = []
    female_ranks = []
    for short_name in outcome_cols:
        res = results.get(short_name, {})
        var_labels.append(ANALYSIS_LABELS.get(short_name, short_name))
        male_ranks.append(res.get("male_mean_rank", 0))
        female_ranks.append(res.get("female_mean_rank", 0))

    x = np.arange(len(var_labels))
    width = 0.35

    ax.bar(x - width / 2, male_ranks, width, label="Male Mean Rank",
           color=male_color, alpha=0.85, edgecolor="white")
    ax.bar(x + width / 2, female_ranks, width, label="Female Mean Rank",
           color=female_color, alpha=0.85, edgecolor="white")

    ax.set_ylabel("Mean Rank", fontsize=11)
    ax.set_title("Male vs Female: Mean Rank Comparison", fontsize=13,
                 fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(var_labels, rotation=15, ha="right")
    ax.legend()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # Add significance stars
    for i, short_name in enumerate(outcome_cols):
        res = results.get(short_name, {})
        if res.get("significant_no_CC"):
            max_h = max(male_ranks[i], female_ranks[i])
            ax.text(i, max_h + 2, "★", ha="center", fontsize=14, color="red")

    plt.tight_layout()
    fpath = os.path.join(output_dir, "fig_mean_rank_comparison.png")
    fig.savefig(fpath, dpi=150, bbox_inches="tight")
    plt.close(fig)
    generated.append(fpath)

    logger.info("Generated %d figures in %s", len(generated), output_dir)
    return generated


# ============================================================
# AUTOMATED TESTS (STEP 13)
# ============================================================
def run_all_assertions(
    results: Dict[str, Dict[str, Any]],
    mt_results: Dict[str, Dict[str, float]],
) -> List[Dict[str, Any]]:
    """Run all automated assertion tests.

    Returns list of test results with pass/fail status.
    """
    tests = []

    for var_name, res in results.items():
        label = ANALYSIS_LABELS.get(var_name, var_name)
        prefix = f"[{label}]"

        # R1 + R2 identity
        tests.append({
            "test": f"{prefix} R1 + R2 = N(N+1)/2",
            "passed": res["rank_sum_identity_ok"],
            "expected": res["expected_rank_sum"],
            "actual": res["R1_plus_R2"],
        })

        # U1 + U2 identity
        tests.append({
            "test": f"{prefix} U1 + U2 = n1 × n2",
            "passed": res["U_identity_ok"],
            "expected": res["n1_times_n2"],
            "actual": res["U1_plus_U2"],
        })

        # N = n1 + n2
        n_ok = res["n1_male"] + res["n2_female"] == res["N"]
        tests.append({
            "test": f"{prefix} N = n1 + n2",
            "passed": n_ok,
            "expected": res["N"],
            "actual": res["n1_male"] + res["n2_female"],
        })

        # Tie counts sum to N
        tests.append({
            "test": f"{prefix} Tie counts sum to N",
            "passed": res["tie_counts_sum_ok"],
            "expected": res["N"],
            "actual": res["tie_counts_sum"],
        })

        # p between 0 and 1
        p_ok = 0 <= res["p_no_CC_manual"] <= 1
        tests.append({
            "test": f"{prefix} 0 ≤ p ≤ 1",
            "passed": p_ok,
            "expected": "[0, 1]",
            "actual": res["p_no_CC_manual"],
        })

        # sigma_U > 0
        sigma_ok = res["sigma_U"] > 0
        tests.append({
            "test": f"{prefix} σ_U > 0",
            "passed": sigma_ok,
            "expected": "> 0",
            "actual": res["sigma_U"],
        })

        # Effect size between -1 and 1
        r_ok = -1 <= res["effect_size_r"] <= 1
        tests.append({
            "test": f"{prefix} -1 ≤ effect size r ≤ 1",
            "passed": r_ok,
            "expected": "[-1, 1]",
            "actual": res["effect_size_r"],
        })

        # SciPy cross-validation (p values should be close)
        if res.get("scipy_p_CC") is not None:
            p_close = abs(res["p_CC_manual"] - res["scipy_p_CC"]) < 0.001
            tests.append({
                "test": f"{prefix} Manual p_CC ≈ SciPy p_CC",
                "passed": p_close,
                "expected": res["scipy_p_CC"],
                "actual": res["p_CC_manual"],
            })

        # SciPy U cross-validation
        if res.get("scipy_U1") is not None:
            u_close = abs(res["U1_male"] - res["scipy_U1"]) < 0.01
            tests.append({
                "test": f"{prefix} Manual U1 ≈ SciPy U1",
                "passed": u_close,
                "expected": res["scipy_U1"],
                "actual": res["U1_male"],
            })

    # Multiple testing: adjusted p >= raw p
    for var_name, mt in mt_results.items():
        label = ANALYSIS_LABELS.get(var_name, var_name)
        prefix = f"[{label}]"

        bonf_ok = mt["bonferroni_p"] >= mt["raw_p"] - 1e-10
        tests.append({
            "test": f"{prefix} Bonferroni p ≥ raw p",
            "passed": bonf_ok,
            "expected": f"≥ {mt['raw_p']:.6f}",
            "actual": mt["bonferroni_p"],
        })

        holm_ok = mt["holm_p"] >= mt["raw_p"] - 1e-10
        tests.append({
            "test": f"{prefix} Holm p ≥ raw p",
            "passed": holm_ok,
            "expected": f"≥ {mt['raw_p']:.6f}",
            "actual": mt["holm_p"],
        })

    return tests


def format_test_results(tests: List[Dict[str, Any]]) -> str:
    """Format test results as readable text."""
    lines = []
    lines.append("=" * 70)
    lines.append("AUTOMATED TEST RESULTS")
    lines.append("=" * 70)

    passed = sum(1 for t in tests if t["passed"])
    total = len(tests)
    lines.append(f"\nTotal: {total} tests | Passed: {passed} | "
                 f"Failed: {total - passed}")
    lines.append("")

    for t in tests:
        status = "✓ PASS" if t["passed"] else "✗ FAIL"
        lines.append(f"  {status}  {t['test']}")
        if not t["passed"]:
            lines.append(f"         Expected: {t['expected']}")
            lines.append(f"         Actual:   {t['actual']}")

    lines.append("")
    if passed == total:
        lines.append("ALL TESTS PASSED ✓")
    else:
        lines.append(f"WARNING: {total - passed} TESTS FAILED ✗")

    return "\n".join(lines)


def export_table_4_4_docx(res: Dict[str, Any], title: str, filepath: str) -> None:
    """Export APA-7 Word table for single Mann-Whitney U test."""
    import docx
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = docx.Document()
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(f"Table: Mann-Whitney U Comparison for {title}")
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(11)
    r_title.font.bold = True

    table = doc.add_table(rows=3, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Group", "N", "Mean Rank", "Mann-Whitney U", "Z", "p-value"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.bold = True

    row1 = [
        "Male",
        str(res["n1_male"]),
        str(res["male_mean_rank"]),
        str(res["U1_male"]),
        str(round(res["Z_no_CC"], 4)),
        f"{res['p_no_CC_manual']:.4f}" if res["p_no_CC_manual"] >= 0.001 else "< 0.001"
    ]
    row2 = [
        "Female",
        str(res["n2_female"]),
        str(res["female_mean_rank"]),
        str(res["U2_female"]),
        "—",
        "—"
    ]

    for idx, r_data in enumerate([row1, row2]):
        for c_idx, val in enumerate(r_data):
            cell = table.rows[idx + 1].cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(9.5)

    def set_cell_border(cell, **kwargs):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/><w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/></w:tcBorders>' % (
            nsdecls('w'),
            kwargs.get('top', 'none'), kwargs.get('top_sz', '4'), kwargs.get('top_color', 'auto'),
            kwargs.get('bottom', 'none'), kwargs.get('bottom_sz', '4'), kwargs.get('bottom_color', 'auto')
        ))
        tcPr.append(tcBorders)

    for cell in table.rows[0].cells:
        set_cell_border(cell, top='single', top_sz='12', bottom='single', bottom_sz='8')
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom='single', bottom_sz='12')

    note_p = doc.add_paragraph()
    r_note = note_p.add_run(
        f"Note. N = {res['N']}. Reported U (min) = {res['U_min']}. Effect size r = {res['effect_size_r']} ({res['effect_size_r_interpretation']}). Decision: {res['decision']}"
    )
    r_note.font.name = "Times New Roman"
    r_note.font.size = Pt(9)
    r_note.font.italic = True

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)

